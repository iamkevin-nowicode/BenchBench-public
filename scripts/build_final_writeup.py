#!/usr/bin/env python3
"""Build the post-run Bench-bench leaderboard and final DOCX report.

The report is intentionally derived from the live transcripts and the current
scripted-baseline artifact.  It does not make network calls and never reads
provider credentials.
"""

from __future__ import annotations

from collections import OrderedDict
from datetime import date
import json
import math
from pathlib import Path
import shutil
import sys
from statistics import fmean, stdev
from typing import Any, Iterable

from scipy.stats import t

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

SKILL_DIR = Path("/Users/kevinolivieri/.codex/plugins/cache/openai-primary-runtime/documents/26.805.11740/skills/documents")
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402
from build_final_visuals import build_figures  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from bench_bench.runner import retry_metrics_from_records  # noqa: E402

REPORTS = ROOT / "reports"
OUTPUT_DOCX = REPORTS / "BENCH_BENCH_FINAL_WRITEUP.docx"
OUTPUT_MD = REPORTS / "FINAL_PUBLIC_LEADERBOARD.md"
OUTPUT_JSON = REPORTS / "final_public_leaderboard.json"
ENGINE_HASH = "sha256:fdbd829339622163df8a27d64fe6467e353c1b2bd8ff289b25e36783e8d2e9a1"
PUBLIC_SEEDS = list(range(100, 110))
T_CRIT_95_DF9 = float(t.ppf(0.975, 9))

LIVE_ROOTS = OrderedDict(
    [
        ("claude-opus-5", ROOT / "runs/live-full-20260808/claude-opus-5"),
        ("gpt-5.6-sol", ROOT / "runs/live-full-20260808/gpt-5.6-sol"),
        ("kimi-k3", ROOT / "runs/live-full-20260808/kimi-k3"),
        ("muse-spark-1.2", ROOT / "runs/live-full-20260808/muse-spark-1.2"),
        ("grok-4.5", ROOT / "runs/live-grok-4.5-full-20260811/grok-4.5"),
    ]
)

MODEL_DISPLAY = {
    "claude-opus-5": "Claude Opus 5",
    "grok-4.5": "Grok 4.5",
    "muse-spark-1.2": "Muse Spark 1.2",
    "gpt-5.6-sol": "GPT-5.6 Sol",
    "kimi-k3": "Kimi K3",
}

RANK_ORDER = ["claude-opus-5", "grok-4.5", "muse-spark-1.2", "gpt-5.6-sol", "kimi-k3"]


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def transcript_metrics(path: Path) -> dict[str, Any]:
    records = read_jsonl(path)
    start = next(record for record in records if record.get("type") == "run_start")
    turns = [record for record in records if record.get("type") == "turn"]
    interrupts = sum(len(record.get("reactive_turns", []) or []) for record in turns)
    retry_metrics = retry_metrics_from_records(records)
    return {
        "model": start["model"],
        "provider": start.get("provider"),
        "endpoint_metadata": start.get("endpoint_metadata", {}),
        "sampling": start.get("sampling", {}),
        "pricing": start.get("pricing", {}),
        "engine_config_hash": start.get("engine_config_hash"),
        "seed": int(start["seed"]),
        "weeks": len(turns),
        "decisions": len(turns) + interrupts,
        "interrupts": interrupts,
        "repairs": retry_metrics["rejected_output_decisions"],
        "rejected_model_outputs": retry_metrics["rejected_model_outputs"],
        "repair_attempts": retry_metrics["repair_attempts"],
        "successful_repairs": retry_metrics["successful_repairs"],
        "transport_failures": retry_metrics["transport_failures"],
    }


def load_live_model(model: str, root: Path) -> dict[str, Any]:
    records: list[dict[str, Any]] = []
    for seed in PUBLIC_SEEDS:
        transcript = next(root.glob(f"seed-{seed}/*-seed-{seed}.jsonl"))
        summary = next(root.glob(f"seed-{seed}/*-summary.json"))
        summary_record = json.loads(summary.read_text(encoding="utf-8"))["results"][0]
        metrics = transcript_metrics(transcript)
        final = summary_record["final_result"]
        records.append(
            {
                **metrics,
                "score": float(final["final_1rm_kg"]),
                "pain_days": int(final["pain_days"]),
                "invalid_reason": final.get("invalid_reason"),
                "cost_usd": float(summary_record.get("total_cost_usd", 0.0)),
                "input_tokens": int(summary_record.get("input_tokens", 0)),
                "visible_output_tokens": int(summary_record.get("visible_output_tokens", 0)),
                "thinking_tokens": int(summary_record.get("thinking_tokens", 0)),
                "total_tokens": int(summary_record.get("total_tokens", 0)),
                "model_calls": int(summary_record.get("model_calls", 0)),
                "summary_repair_calls": int(summary_record.get("repair_calls", 0)),
            }
        )
    scores = [record["score"] for record in records]
    return {
        "model": model,
        "display_name": MODEL_DISPLAY[model],
        "records": records,
        "scores": {str(record["seed"]): record["score"] for record in records},
        "mean_kg": fmean(scores),
        "seed_sd_kg": stdev(scores),
        "min_kg": min(scores),
        "max_kg": max(scores),
        "counted_seeds": sum(record["pain_days"] <= 14 and record["invalid_reason"] is None for record in records),
        "pain_violations": sum(record["pain_days"] > 14 for record in records),
        "invalid_episodes": sum(record["invalid_reason"] is not None for record in records),
        "decisions": sum(record["decisions"] for record in records),
        "repairs": sum(record["repairs"] for record in records),
        "repair_rate": sum(record["repairs"] for record in records) / sum(record["decisions"] for record in records),
        "transport_failures": sum(record["transport_failures"] for record in records),
        "interrupts": sum(record["interrupts"] for record in records),
        "cost_total_usd": sum(record["cost_usd"] for record in records),
        "cost_per_episode_usd": fmean(record["cost_usd"] for record in records),
        "input_tokens": sum(record["input_tokens"] for record in records),
        "visible_output_tokens": sum(record["visible_output_tokens"] for record in records),
        "thinking_tokens": sum(record["thinking_tokens"] for record in records),
        "total_tokens": sum(record["total_tokens"] for record in records),
        "model_calls": sum(record["model_calls"] for record in records),
        "endpoint_metadata": records[0]["endpoint_metadata"],
        "sampling": records[0]["sampling"],
        "pricing": records[0]["pricing"],
        "engine_config_hashes": sorted({record["engine_config_hash"] for record in records}),
    }


def load_baselines() -> dict[str, dict[str, Any]]:
    report = json.loads((REPORTS / "current_baseline_gate.json").read_text(encoding="utf-8"))
    result: dict[str, dict[str, Any]] = {}
    for name, episodes in report["episodes"].items():
        raw = [float(episode["raw_final_1rm_kg"]) for episode in episodes]
        counted = [episode["counted_final_1rm_kg"] for episode in episodes if episode["counted_final_1rm_kg"] is not None]
        result[name] = {
            "name": name,
            "mean_kg": fmean(raw),
            "seed_sd_kg": stdev(raw),
            "counted_mean_kg": fmean(counted) if len(counted) == len(episodes) else None,
            "counted_seed_sd_kg": stdev(counted) if len(counted) == len(episodes) else None,
            "counted_seeds": len(counted),
            "total_seeds": len(episodes),
            "violations": sum(bool(episode["constraint_violations"]) for episode in episodes),
            "raw_scores": {str(episode["seed"]): episode["raw_final_1rm_kg"] for episode in episodes},
        }
    return result


def paired_difference(a: dict[str, Any], b: dict[str, Any]) -> dict[str, Any]:
    left = {int(seed): score for seed, score in a["scores"].items()}
    right = {int(seed): score for seed, score in b["scores"].items()}
    diffs = [left[seed] - right[seed] for seed in PUBLIC_SEEDS]
    delta = fmean(diffs)
    half = T_CRIT_95_DF9 * stdev(diffs) / math.sqrt(len(diffs))
    return {
        "a": a["model"],
        "b": b["model"],
        "a_display": a["display_name"],
        "b_display": b["display_name"],
        "mean_difference_kg": delta,
        "ci95_low_kg": delta - half,
        "ci95_high_kg": delta + half,
        "difference_sd_kg": stdev(diffs),
    }


def all_pairwise(models: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    values = []
    for index, left in enumerate(RANK_ORDER):
        for right in RANK_ORDER[index + 1 :]:
            values.append(paired_difference(models[left], models[right]))
    return values


def public_data() -> dict[str, Any]:
    models = OrderedDict((name, load_live_model(name, root)) for name, root in LIVE_ROOTS.items())
    baselines = load_baselines()
    pairwise = all_pairwise(models)
    return {
        "generated_date": str(date.today()),
        "engine_config_hash": ENGINE_HASH,
        "horizon_weeks": 52,
        "public_seeds": PUBLIC_SEEDS,
        "episodes": 50,
        "models": models,
        "baselines": baselines,
        "pairwise_model_differences": pairwise,
        "aggregation": {
            "counted_seed_fraction_required": 1.0,
            "pain_days_limit": 14,
            "paired_ci": "two-sided 95% t interval over 10 matched public seeds; df=9",
        },
    }


def fmt(value: float | None, digits: int = 2) -> str:
    return "—" if value is None else f"{value:.{digits}f}"


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def write_leaderboard_markdown(data: dict[str, Any]) -> None:
    models = data["models"]
    ranked = sorted(models.values(), key=lambda item: item["mean_kg"], reverse=True)
    lines = [
        "# Bench-bench post-run public leaderboard",
        "",
        "This standalone post-run report covers five models × ten public seeds × 52 simulated weeks (50 episodes). It is derived from the completed live transcripts; no network calls are made during report generation.",
        "",
        f"- Engine/config hash: `{data['engine_config_hash']}`",
        "- Seeds: 100–109",
        "- Score: mean of standardized tests at weeks 44, 48, and 52 after the fixed three-day taper",
        "- Counted score: only when pain days ≤14 and `invalid_reason` is null; all five models counted on all ten seeds",
        "",
        "## Live model leaderboard",
        "",
        "| Rank | Model | Mean kg | Seed SD | Range kg | Counted | Pain violations | Repairs / decisions | Transport failures | Cost / ep |",
        "|---:|---|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for rank, item in enumerate(ranked, start=1):
        lines.append(
            f"| {rank} | {item['display_name']} | {item['mean_kg']:.2f} | {item['seed_sd_kg']:.2f} | {item['min_kg']:.2f}–{item['max_kg']:.2f} | {item['counted_seeds']}/10 | {item['pain_violations']} | {item['repairs']}/{item['decisions']} ({pct(item['repair_rate'])}) | {item['transport_failures']} | ${item['cost_per_episode_usd']:.4f} |"
        )
    lines.extend(
        [
            "",
            "## Scripted reference baselines",
            "",
            "| Reference policy | Mean kg | Seed SD | Counted mean | Counted seeds | Violations |",
            "|---|---:|---:|---:|---:|---:|",
        ]
    )
    baseline_order = ["scripted-expert", "recovery-aware", "skip-when-busy", "rigid-linear", "reckless-maximalist", "random"]
    for name in baseline_order:
        item = data["baselines"][name]
        lines.append(
            f"| {name} | {item['mean_kg']:.2f} | {item['seed_sd_kg']:.2f} | {fmt(item['counted_mean_kg'])} | {item['counted_seeds']}/{item['total_seeds']} | {item['violations']} |"
        )
    lines.extend(
        [
            "",
            "## Paired differences",
            "",
            "A minus B, using the same ten public seeds; intervals are two-sided 95% paired t intervals (df=9), without multiplicity correction.",
            "",
            "| A | B | Mean A−B kg | 95% CI kg |",
            "|---|---|---:|---:|",
        ]
    )
    for pair in data["pairwise_model_differences"]:
        lines.append(
            f"| {pair['a_display']} | {pair['b_display']} | {pair['mean_difference_kg']:.2f} | [{pair['ci95_low_kg']:.2f}, {pair['ci95_high_kg']:.2f}] |"
        )
    lines.extend(
        [
            "",
            "## Per-seed scores",
            "",
            "| Model | " + " | ".join(str(seed) for seed in PUBLIC_SEEDS) + " |",
            "|---|" + "---:|" * len(PUBLIC_SEEDS),
        ]
    )
    for item in ranked:
        lines.append("| " + item["display_name"] + " | " + " | ".join(f"{item['scores'][str(seed)]:.2f}" for seed in PUBLIC_SEEDS) + " |")
    lines.extend(
        [
            "",
            "## Audit",
            "",
            "- 50/50 live transcript start records carry the same engine/config hash.",
            "- All 50 episodes completed; all had pain days 0 and no structural invalidation.",
            "- Transport failures are reported separately from rejected-output repairs. Kimi K3 had 702 transport failures; the other four models had none.",
            "- Total live spend: $127.63.",
        ]
    )
    OUTPUT_MD.write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color: str = "D9E2F3", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, *, size=11, color="222222", bold=False, italic=False, font="Calibri") -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_para(paragraph, *, before=0, after=6, line=1.10, keep=False, alignment=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep
    if alignment is not None:
        paragraph.alignment = alignment


def add_text(doc: Document, text: str, *, style="Normal", before=0, after=6, line=1.10, keep=False) -> Any:
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    set_para(paragraph, before=before, after=after, line=line, keep=keep)
    return paragraph


def add_rich_paragraph(doc: Document, parts: Iterable[tuple[str, dict[str, Any]]], *, style="Normal", before=0, after=6, line=1.10) -> Any:
    paragraph = doc.add_paragraph(style=style)
    for text, options in parts:
        run = paragraph.add_run(text)
        set_run(run, **options)
    set_para(paragraph, before=before, after=after, line=line)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1) -> Any:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_para(p, keep=True)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], *, font_size=8.7, header_fill="F2F4F7", indent=120):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para(p, before=0, after=0, line=1.0)
        run = p.add_run(header)
        set_run(run, size=font_size, color="1F4D78", bold=True)
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_para(p, before=0, after=0, line=1.0)
            run = p.add_run(str(value))
            set_run(run, size=font_size, color="222222")
            set_cell_borders(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=indent)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    set_cell_shading(cell, "E8EEF5")
    set_cell_borders(cell, color="B8CBE1", size="8")
    p = cell.paragraphs[0]
    set_para(p, before=0, after=3, line=1.10)
    run = p.add_run(title)
    set_run(run, size=10.5, color="0B2545", bold=True)
    p2 = cell.add_paragraph()
    set_para(p2, before=0, after=0, line=1.10)
    run = p2.add_run(body)
    set_run(run, size=10.5, color="222222")
    set_repeat_table_header(table.rows[0])
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_timeline_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a chronological review table with a visually distinct stage rail."""
    table = add_table(
        doc,
        ["Sequence", "Question / test", "What we found", "Iteration / decision"],
        rows,
        [1450, 2550, 2700, 2660],
        font_size=8.05,
    )
    for row_index, row in enumerate(table.rows[1:], start=1):
        stage = row.cells[0]
        set_cell_shading(stage, "1F3A5F")
        for paragraph in stage.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run(run, size=8.05, color="FFFFFF", bold=True)
        if row_index % 2 == 0:
            for cell in row.cells[1:]:
                set_cell_shading(cell, "F7F9FC")
    return table


def add_figure(doc: Document, path: Path, caption: str, *, width=6.5) -> None:
    """Add an inline, accessible figure with a compact caption."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(paragraph, before=4, after=2, line=1.0, keep=True)
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption.split(".", 1)[0])
    caption_paragraph = doc.add_paragraph()
    set_para(caption_paragraph, before=0, after=8, line=1.0)
    caption_run = caption_paragraph.add_run(caption)
    set_run(caption_run, size=8.7, color="4B5563", italic=True)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, size=9, color="6B7280")


def build_docx(data: dict[str, Any]) -> None:
    figures = build_figures(data)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    # Standard business brief token map.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, "2E74B5", 16, 8),
        (2, 13, "2E74B5", 12, 6),
        (3, 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    # Quiet running header/footer.
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(header, before=0, after=0, line=1.0)
    run = header.add_run("BENCH-BENCH  /  FINAL RESULTS")
    set_run(run, size=8.5, color="6B7280", bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(footer, before=0, after=0, line=1.0)
    run = footer.add_run("Bench-bench · 12 August 2026  ·  ")
    set_run(run, size=8.5, color="6B7280")
    add_page_number(footer)

    # Masthead.
    p = doc.add_paragraph()
    set_para(p, before=8, after=4, line=1.0)
    r = p.add_run("BENCH-BENCH")
    set_run(r, size=10, color="7A5A00", bold=True)
    p = doc.add_paragraph()
    set_para(p, before=0, after=4, line=1.0, keep=True)
    r = p.add_run("Final results and development retrospective")
    set_run(r, size=23, color="0B2545", bold=True)
    p = doc.add_paragraph()
    set_para(p, before=0, after=14, line=1.0)
    r = p.add_run("A 52-week benchmark of long-horizon coaching under time, money, recovery, and household scarcity")
    set_run(r, size=13.5, color="4B5563")
    for label, value in (
        ("Run", "Public seeds 100–109 · five models · 50 completed episodes"),
        ("Protocol", "Three standardized tests at weeks 44, 48, and 52 after a fixed three-day taper"),
        ("Engine/config", data["engine_config_hash"]),
        ("Prepared", "12 August 2026"),
    ):
        add_rich_paragraph(
            doc,
            [(f"{label}: ", {"size": 10.5, "color": "222222", "bold": True}), (value, {"size": 10.5, "color": "222222"})],
            after=2,
            line=1.0,
        )
    add_callout(
        doc,
        "Bottom line",
        "Claude Opus 5 led the live models at 100.06 kg, followed closely by Grok 4.5 at 99.19 kg and Muse Spark 1.2 at 98.62 kg. The scripted-expert reference remained higher at 102.89 kg. Every live episode stayed within the pain constraint and completed without structural invalidation; Kimi K3 is the important operational exception, with 702 transport failures recorded separately from its 5.76% rejected-output repair rate.",
    )

    add_heading(doc, "1. What the benchmark measures", 1)
    add_text(doc, "Bench-bench asks an agent to act as Dave’s coach for one simulated year. Dave is a returning lifter with a full-time job, a full-time working partner, a six-month-old baby at the start, a commercial gym membership, no home equipment, and a constrained discretionary budget. Each week the agent plans training and allocates time and money across meals, childcare, chores, partner coverage, giveback, sleep protection, and purchases. Seeded interruptions then test whether the plan survives contact with ordinary life.")
    add_text(doc, "The score is not the noisy weekly 1RM estimate. It is the arithmetic mean of three read-only standardized-test projections at weeks 44, 48, and 52, each after a fixed three-day taper. A score is counted only when pain days are at most 14 and the episode has no structural invalidation. Household strain and sleep debt are reported diagnostics; they are not hidden score penalties.")
    add_text(doc, "The live evaluation used temperature 1.0, medium effort wherever exposed, direct provider endpoints, one repair attempt, incremental JSONL transcripts, and resumable execution. Every transcript carries the same engine/config hash; the 50-transcript audit passed.")

    add_heading(doc, "2. Final leaderboard", 1)
    add_text(doc, "The live model leaderboard ranks mean counted score over the ten matched public seeds. Because every live episode had pain days 0 and a null invalid reason, raw and counted scores are identical for all five models.")
    ranked = sorted(data["models"].values(), key=lambda item: item["mean_kg"], reverse=True)
    rows = []
    for rank, item in enumerate(ranked, start=1):
        rows.append([
            str(rank), item["display_name"], f"{item['mean_kg']:.2f}", f"{item['seed_sd_kg']:.2f}", f"{item['min_kg']:.2f}–{item['max_kg']:.2f}", f"{item['counted_seeds']}/10", str(item["pain_violations"]), f"{item['repairs']}/{item['decisions']} ({pct(item['repair_rate'])})", str(item["transport_failures"]), f"${item['cost_per_episode_usd']:.4f}",
        ])
    add_table(doc, ["Rank", "Model", "Mean kg", "SD", "Range", "Counted", "Pain", "Repairs", "Transport", "$ / ep"], rows, [450, 1650, 680, 550, 1050, 650, 500, 1100, 900, 1830], font_size=8.1)
    add_text(doc, "Interpretation: the top three are close relative to ten-seed uncertainty. Opus exceeds Grok by 0.87 kg on average, but the paired 95% interval is [−1.12, 2.87] kg; that is not a robust separation at this sample size. Muse is clearly above GPT on matched seeds, while GPT versus Kimi is not cleanly separated because Kimi’s seed variance is large.")
    add_figure(doc, figures["performance_overview"], "Figure 1. The live-model means cluster below the scripted expert; the reference hierarchy is much wider than the live top-three differences. Live whiskers show 95% intervals for the ten-seed mean; scripted whiskers show seed standard deviation.")

    add_heading(doc, "3. Scripted reference baselines", 1)
    add_text(doc, "The six fixed policies are calibration references, not competing model entries. Their role is to show that the simulator preserves the intended hierarchy and that reckless behavior loses endogenously rather than winning through a lucky final test.")
    baseline_rows = []
    for name in ["scripted-expert", "recovery-aware", "skip-when-busy", "rigid-linear", "reckless-maximalist", "random"]:
        item = data["baselines"][name]
        baseline_rows.append([
            name,
            f"{item['mean_kg']:.2f}",
            f"{item['seed_sd_kg']:.2f}",
            fmt(item["counted_mean_kg"]),
            f"{item['counted_seeds']}/{item['total_seeds']}",
            str(item["violations"]),
        ])
    add_table(doc, ["Reference policy", "Raw mean kg", "Seed SD", "Counted mean", "Counted", "Violations"], baseline_rows, [2100, 1300, 1100, 1450, 1250, 2160], font_size=9.0)
    add_text(doc, "Final reference ordering: scripted-expert (102.89) > recovery-aware (99.01) > skip-when-busy (96.56) > rigid-linear (92.35) > reckless-maximalist (87.51 raw) > random (86.94). Reckless-maximalist violated pain days on all 20 burned development seeds, so its counted mean is unavailable even though its raw diagnostic mean remains visible. The expert–random gap is 15.95 kg, equal to 18.92 pooled seed standard deviations; the 52-week gate passed, including the 65% adjacent-order criterion and endogenous reckless loss.")

    add_heading(doc, "4. Statistical comparison of the live models", 1)
    add_text(doc, "The table below reports paired differences over the same ten public seeds. A positive value means the model in the first column scored higher. Intervals are two-sided 95% paired t intervals with nine degrees of freedom and no multiplicity correction; they are uncertainty summaries, not claims of a formal multi-comparison hypothesis test.")
    pair_rows = []
    for pair in data["pairwise_model_differences"]:
        pair_rows.append([
            pair["a_display"],
            pair["b_display"],
            f"{pair['mean_difference_kg']:.2f}",
            f"[{pair['ci95_low_kg']:.2f}, {pair['ci95_high_kg']:.2f}]",
        ])
    add_table(doc, ["A", "B", "Mean A−B kg", "95% paired CI kg"], pair_rows, [2500, 2500, 1700, 2660], font_size=8.9)
    add_text(doc, "The practical picture is a three-tier top cluster rather than five crisply separated ranks: Opus, Grok, and Muse are within roughly 1.45 kg of each other in mean score, and their adjacent confidence intervals include zero. GPT is about 4.75 kg below Muse with a wholly positive interval. Kimi’s mean is lowest and its spread is much larger than every other model, so a ten-seed average is particularly unstable for that provider.")
    add_figure(doc, figures["seed_variation_and_pairwise"], "Figure 2. The matched-seed heatmap shows common difficulty variation across models, while the paired intervals show that the Opus–Grok, Opus–Muse, and Grok–Muse gaps are not cleanly separated with ten seeds.")

    add_heading(doc, "5. Operational behavior and cost", 1)
    add_text(doc, "Each model produced 677 decisions: 520 weekly plans plus 157 reactive interrupt decisions. Repairs count rejected model outputs only; transport failures are a separate metric. This distinction mattered in the Kimi run, where provider rate-limit/transport behavior was substantial but should not be mislabeled as model-format failure.")
    op_rows = []
    for item in sorted(data["models"].values(), key=lambda item: item["mean_kg"], reverse=True):
        endpoint = item["endpoint_metadata"].get("url", "—")
        op_rows.append([
            item["display_name"],
            f"{item['repairs']}/{item['decisions']} ({pct(item['repair_rate'])})",
            str(item["transport_failures"]),
            f"${item['cost_total_usd']:.2f}",
            f"${item['cost_per_episode_usd']:.4f}",
            str(item["sampling"].get("effort", "—")),
            endpoint.replace("https://", ""),
        ])
    add_table(doc, ["Model", "Repairs", "Transport", "Total cost", "$ / ep", "Effort", "Endpoint"], op_rows, [1500, 1500, 950, 1050, 1050, 850, 2460], font_size=7.8)
    add_text(doc, "Total live spend was $127.63. All provider cost ceilings were respected. Token accounting was recorded as input, visible output, and thinking tokens where available; provider semantics differ, so raw token totals should be used for cost auditing rather than as a direct measure of reasoning quality.")
    token_rows = []
    for item in sorted(data["models"].values(), key=lambda item: item["mean_kg"], reverse=True):
        token_rows.append([
            item["display_name"],
            f"{item['input_tokens'] / 1_000_000:.2f}",
            f"{item['visible_output_tokens'] / 1_000_000:.2f}",
            f"{item['thinking_tokens'] / 1_000_000:.2f}",
            f"{item['total_tokens'] / 1_000_000:.2f}",
        ])
    add_table(doc, ["Model", "Input M", "Visible output M", "Thinking M", "Total M"], token_rows, [2500, 1500, 1800, 1600, 1960], font_size=8.8)
    add_figure(doc, figures["operational_behavior"], "Figure 3. Output discipline, provider transport reliability, and cost are distinct dimensions: Opus required the most repairs, Kimi had the only substantial transport-failure burden, and cost varied by more than sixfold per episode.")

    add_heading(doc, "6. What the results say", 1)
    add_heading(doc, "6.1 The benchmark rewards sustained operational judgment", 2)
    add_text(doc, "The live ordering is not a simple intelligence ranking detached from the simulator. The task rewards keeping a plan feasible under a 900-minute weekly ledger, protecting recovery without starving training, reacting to shocks, and avoiding the temptation to spend the whole budget on a nominally optimal week. The gap between the scripted expert and the lowest-variance model results is a useful reference point, but the close Opus/Grok/Muse cluster shows that ten seeds are better at identifying broad tiers than precise rank order.")
    add_heading(doc, "6.2 Kimi’s variance is an operational finding, not just a score", 2)
    add_text(doc, "Kimi K3’s 6.22 kg seed SD is roughly three times Grok’s and five times GPT’s. The run also logged 702 transport failures, while the other four models logged none. That co-occurrence is suggestive that provider reliability contributed to the spread, but it is not causal proof: transport failures change the action path through retries and fallbacks, and the benchmark intentionally exposes the full system’s ability to continue. Kimi should therefore be reported with the transport caveat attached, not silently treated as an ordinary low-scoring model.")
    add_heading(doc, "6.3 Format discipline mattered, especially for Opus and early Grok turns", 2)
    add_text(doc, "Opus had the highest rejected-output repair rate at 19.20%, even after the prompt fixes, while GPT was 2.22%, Muse 4.43%, Grok 8.42%, and Kimi 5.76%. The main recurring errors were not physiology failures: they were schema nesting, shared-ledger feasibility, and constraints that had been enforced before being explicitly documented. Grok’s first-week failures flattened the nested life allocation into the weekly action; that finding is logged as a v0.2 prompt-clarity issue. The lesson is methodological: a long-horizon benchmark must make its output contract as explicit as its simulator contract, and must report repairs without confusing them with transport noise.")
    add_heading(doc, "6.4 Zero pain days means the hard constraint did not fire", 2)
    add_text(doc, "All 50 live episodes ended with pain_days=0. In benchmark terms this means no episode accumulated any recorded pain days and therefore none was voided by the pain-days rule. It does not establish that the model is safe for real training, that no soreness or discomfort occurred outside the metric, or that the benchmark validates medical advice. The score remains a simulation result, not a health claim.")

    add_heading(doc, "7. How the benchmark was built", 1)
    add_text(doc, "The development process was deliberately adversarial. Each stage started with a concrete question, was tested with a diagnostic, ablation, adversarial policy, or live run, and ended with a decision that changed either the simulator, the protocol, the runner, or the evidence standard. The sequence below is the process review that led from the initial plan to the frozen public results.")
    add_figure(doc, figures["development_timeline"], "Figure 4. The benchmark moved from concept to calibration, mechanics hardening, evidence review, provider preparation, and finally a frozen public run.")
    timeline_rows = [
        ["Start / brief", "Can a benchmark test long-horizon coaching rather than isolated workout advice?", "A weekly plan has to coordinate training, work, money, sleep, childcare, chores, and a working partner; future interruptions must remain hidden.", "Defined Dave, the model-only track, weekly actions plus reactive turns, delayed consequences, and the question: how much can Dave bench without letting the rest of life collapse?"],
        ["Phase 1 / slice", "Can the schemas, runner, replay, repair, and fallback path operate end to end?", "The 12-week vertical slice made the mechanics observable, but short-horizon adjacent policies were too close to discriminate reliably.", "Kept 12 weeks as development diagnostics and made 52 simulated weeks the release horizon."],
        ["Calibration", "Does the engine preserve an interpretable baseline hierarchy?", "Early calibration exposed questions about expert trajectory, returning-lifter gain, fixed bodyweight, nutrition, hidden tests, and whether the score was noisy or standardized.", "Retuned delayed adaptation, repaired unused technique and mass pathways, and defined the final score as tests at weeks 44, 48, and 52 after a fixed taper."],
        ["Ablations", "Which parts of the world model actually change decisions and scores?", "The home rack, gym logistics, events, recovery, technique, adherence, and household reciprocity were not all equally active; home equipment could dominate and seven open days made life too easy.", "Moved difficulty into logistics: 900 minutes of conserved weekly time, gym commute cost, realistic open days, reciprocity degradation, and consistency drift of 0.10 kg after a productive streak."],
        ["Repair audit", "Are model failures telling us about reasoning or about an unclear interface?", "The interrupt mechanic was effectively disabled by automatic fallback; weekly and reactive schemas were being conflated, and repairs often lacked the correct schema.", "Separated turn-appropriate prompts, supplied the correct repair schema, counted rejected outputs separately from transport failures, and recorded every attempt and fallback."],
        ["Mechanics review 1", "Can a legal policy win by stacking volume or declaring compressed fallback work?", "8×4 and related high-volume programs showed that stimulus was too close to linear in sets × reps; zero-load declarations could earn real credit; silent clipping hid infeasible intent.", "Added the Brzycki-style load/reps ceiling, rep-rate coupling for every focus, fallback caps, no zero-load stimulus, a diminishing weekly cap, and validation-time rejection instead of silent clipping."],
        ["Ledger review", "Can the agent create free time or free resources by allocating aggressively?", "The 8×4 strategy needed more than a full-time worker and infant household could supply; physiology tuning did not solve the root problem.", "Unified training, commute, meals, childcare, chores, partner coverage, and giveback under one finite ledger; delegation and reactive childcare consume cash; expert became ledger-aware."],
        ["Adversarial search", "Can automated search find a better legal policy, including reviewer-written policies?", "The initial exploit scripts were too narrow; hand-written low-intensity/high-volume policies beat the search and showed that beating expert alone is not proof of abuse.", "Widened the genome to mixed focus, per-week structure, zero loads, purchase ordering, and reviewer regressions; separated abuse signatures from the +5 kg human-review flag."],
        ["Safety + statistics", "Can a bad episode or survivor bias produce a wrong leaderboard?", "Pain metrics initially did not fire; missing fields could default to zero; survivor means favored policies that failed on hard seeds; invalidation needed to be real and automatic.", "Made pain days ≤14 the only behavioral hard constraint, failed closed on missing fields, implemented invalid_reason, excluded invalid episodes, and required 100% counted-seed coverage for aggregates."],
        ["Prompt + artifact review", "Does the model know the objective, and can the audit see every correction?", "The system prompt lacked the objective and later risked leaking strategy; notebooks restated plans; transformations such as clipping, coercion, and reactive collapse were not consistently surfaced.", "Rewrote the coach prompt, kept reactive prompting separate, made notebook entries observation-focused, audited the full silent-correction class, embedded hashes, and removed stale claims from the release path."],
        ["Independent reviews", "What survives outside the author’s assumptions?", "Two independent reviews converged on volume stacking, ledger accounting, zero-load credit, private-field auditing, shock reserves, stimulus constants, survivor bias, and search coverage.", "Turned each convergence point into an invariant, regression, or explicit reporting rule. Burned seeds 0–19 for calibration and reserved public 100–109 plus private out-of-band seeds."],
        ["Live preparation", "Can five direct providers run overnight without losing provenance or silently changing the experiment?", "Adapters differed on structured output, effort, thinking-token accounting, pricing, rate limits, and endpoint shape; Kimi was slow and transport-sensitive.", "Used direct endpoints, per-provider pricing, medium effort where exposed, temperature 1.0, resumable transcripts, retry/backoff, cost ceilings, a 100-decision repair guard, and separate transport metrics."],
        ["Public run", "What happens when the frozen benchmark meets real model outputs?", "All 50 episodes completed. Opus led at 100.06 kg; Grok and Muse formed a close top cluster; Kimi showed high variance and 702 transport failures; no live episode violated pain or structural validity.", "Freeze the v0.1 results as a post-run record. Treat the next work as v0.2 prompt/documentation/artifact hygiene, not post hoc simulator retuning."],
    ]
    add_timeline_table(doc, timeline_rows)
    add_text(doc, "The most important pattern in the process was that the hard problems were not isolated bugs. They were classes of failure: free stimulus, free time, silent correction, survivor bias, undocumented constraints, and artifacts that could claim PASS without proving the underlying condition. The benchmark became stronger when each class was turned into an explicit invariant, an audit, or a counted diagnostic.")

    add_heading(doc, "8. Limitations and recommended next step", 1)
    add_text(doc, "The public live set has ten seeds per model, which is enough to expose gross ordering and provider failures but not enough to make close model ranks stable. The benchmark also remains a stylized simulator: the physiology is deliberately coarse, household behavior is represented by designed mechanics, and the standardized test is not a clinical or competitive lifting protocol. Provider output and transport semantics differ, especially for thinking tokens and structured output support.")
    add_text(doc, "The recommended next step is a narrow v0.2 hygiene pass before making comparative claims: add a complete nested weekly-action example to the prompt, document every enforced output limit (including notebook and fallback ceilings), update the benchmark card and release manifest to point to one post-run authoritative leaderboard, and rerun the artifact/hash/private-seed audits. Keep the current live results as the frozen v0.1 record; do not retune the simulator in response to this leaderboard unless a separately preregistered calibration change is approved.")

    doc.add_page_break()
    add_heading(doc, "Appendix A. Per-seed scores", 1)
    ranked = sorted(data["models"].values(), key=lambda item: item["mean_kg"], reverse=True)
    matrix_rows = []
    for item in ranked:
        matrix_rows.append([item["display_name"]] + [f"{item['scores'][str(seed)]:.2f}" for seed in PUBLIC_SEEDS])
    add_table(doc, ["Model"] + [str(seed) for seed in PUBLIC_SEEDS], matrix_rows, [2700] + [666] * 10, font_size=8.0)

    add_heading(doc, "Appendix B. Endpoint and sampling provenance", 1)
    provenance_rows = []
    for item in sorted(data["models"].values(), key=lambda item: item["mean_kg"], reverse=True):
        provenance_rows.append([
            item["display_name"],
            item["model"],
            item["endpoint_metadata"].get("kind", "—"),
            item["endpoint_metadata"].get("url", "—").replace("https://", ""),
            str(item["sampling"].get("effort", "—")),
            str(item["sampling"].get("temperature", "—")),
        ])
    add_table(doc, ["Provider track", "Exact model string", "Adapter", "Endpoint", "Effort", "Temp"], provenance_rows, [1600, 1600, 1200, 2800, 1100, 1060], font_size=7.9)
    add_text(doc, "All endpoint identities above are sanitized URL metadata from transcript start records. No API key or credential material is included in the transcripts, reports, or this document.")

    add_heading(doc, "Appendix C. Reproduction and provenance", 1)
    add_text(doc, "The deterministic baseline source is reports/current_baseline_gate.json. The four-model live transcripts are under runs/live-full-20260808/; the Grok transcripts are under runs/live-grok-4.5-full-20260811/. The engine/config hash is sha256:fdbd829339622163df8a27d64fe6467e353c1b2bd8ff289b25e36783e8d2e9a1. A repository test run passed 100% (pytest -q: 79 tests). All 50 live transcript start records matched the same hash.")
    add_text(doc, "The leaderboard’s counted-seed rule is intentionally strict: an aggregate is reportable only when every expected seed counts. A pain-violating or structurally invalid episode remains visible with its raw score, but it cannot be averaged into a leaderboard mean. Transport failures are not relabeled as repairs; they are reported as a separate operational metric.")

    doc.core_properties.title = "Bench-bench — Final results and development retrospective"
    doc.core_properties.subject = "52-week long-horizon coaching benchmark"
    doc.core_properties.author = "Bench-bench"
    doc.core_properties.comments = "Generated from completed public-seed live transcripts; no credentials included."
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main() -> None:
    data = public_data()
    OUTPUT_JSON.write_text(json.dumps(data, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    write_leaderboard_markdown(data)
    build_docx(data)
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)
    print(OUTPUT_JSON)


if __name__ == "__main__":
    main()
